import os
import pandas as pd
import fitz
import tenseal as ts
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from docx import Document
from io import BytesIO
from PIL import Image
import pytesseract
import openpyxl
from openpyxl.styles import PatternFill
from openpyxl.comments import Comment
import json

# --- Initialisation globale ---
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

def setup_context(context_path="tenseal_context.ctx"):
    context = ts.context(
        ts.SCHEME_TYPE.BFV,
        poly_modulus_degree=4096,
        plain_modulus=1032193
    )
    context.generate_galois_keys()
    context.generate_relin_keys()
    with open(context_path, "wb") as f:
        f.write(context.serialize(save_public_key=True, save_secret_key=True, save_galois_keys=True, save_relin_keys=True))
    return context

def get_ctx(context_path="tenseal_context.ctx"):
    if not os.path.exists(context_path):
        return setup_context(context_path)
    with open(context_path, "rb") as f:
        return ts.context_from(f.read())

def detect_file_type(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext in ['.pdf', '.txt', '.docx', '.csv', '.xlsx']:
        return ext[1:]
    else:
        return 'unknown'

# ------------- Extraction, Détection, Masquage -------------
def extract_text_from_csv(csv_path: str) -> str:
    encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    for encoding in encodings_to_try:
        try:
            df = pd.read_csv(csv_path, encoding=encoding)
            return df.to_string(index=False)
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"Erreur lors de l'extraction CSV avec {encoding}: {e}")
            continue
    return ""

def detect_sensitive_csv(csv_path: str, language: str = 'en'):
    text_content = extract_text_from_csv(csv_path)
    if not text_content:
        return {'0': []}
    analysis = analyzer.analyze(text=text_content, language=language)
    sensitive_items = []
    for result in analysis:
        sensitive_items.append({
            'entity_type': result.entity_type,
            'score': result.score,
            'text': text_content[result.start:result.end],
            'start': result.start,
            'end': result.end
        })
    return {'0': sensitive_items}

def mask_csv_file(csv_path: str, output_path: str, language: str = 'en'):
    encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    df = None
    for encoding in encodings_to_try:
        try:
            df = pd.read_csv(csv_path, encoding=encoding)
            break
        except UnicodeDecodeError:
            continue
        except Exception:
            continue
    if df is None:
        return
    masked_df = df.copy()
    for column in df.columns:
        for index, value in df[column].items():
            if pd.isna(value):
                continue
            text_value = str(value)
            analysis = analyzer.analyze(text=text_value, language=language)
            if analysis:
                masked_result = anonymizer.anonymize(text=text_value, analyzer_results=analysis)
                masked_df.at[index, column] = masked_result.text
    masked_df.to_csv(output_path, index=False, encoding='utf-8')

def process_csv_file(csv_path: str, output_dir: str = "output"):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    text_content = extract_text_from_csv(csv_path)
    if not text_content:
        return {'0': []}, None, None
    analysis = analyzer.analyze(text=text_content, language='en')
    sensitive_items = []
    for result in analysis:
        sensitive_items.append({
            'entity_type': result.entity_type,
            'score': result.score,
            'text': text_content[result.start:result.end],
            'start': result.start,
            'end': result.end
        })
    detected = {'0': sensitive_items}
    masked_output = os.path.join(output_dir, f"masked_{os.path.basename(csv_path)}")
    mask_csv_file(csv_path, masked_output)
    encrypted_output = os.path.join(output_dir, f"encrypted_{os.path.basename(csv_path)}")
    encrypt_sensitive_data(csv_path, encrypted_output)
    return detected, masked_output, encrypted_output

# --- PDF ---
def extract_text_from_pdf(pdf_path: str, use_ocr: bool = True) -> dict:
    doc = fitz.open(pdf_path)
    text_data = {}
    for page_number in range(len(doc)):
        page = doc[page_number]
        text = page.get_text().strip()
        if not text and use_ocr:
            pix = page.get_pixmap(dpi=300)
            img = Image.open(BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(img, lang='eng+fra')
        text_data[page_number] = text
    doc.close()
    return text_data

def detect_sensitive_pdf(pdf_path: str, language: str = 'en'):
    text_data = extract_text_from_pdf(pdf_path)
    doc = fitz.open(pdf_path)
    results = {}
    for page_num, text in text_data.items():
        if not text.strip():
            continue
        page = doc[int(page_num)]
        analysis = analyzer.analyze(text=text, language=language)
        sensitive_items = []
        for result in analysis:
            matched_text = text[result.start:result.end]
            bboxes = page.search_for(matched_text)
            for bbox in bboxes:
                sensitive_items.append({
                    'entity_type': result.entity_type,
                    'score': result.score,
                    'text': matched_text,
                    'bbox': [bbox.x0, bbox.y0, bbox.x1, bbox.y1]
                })
        results[str(page_num)] = sensitive_items
    doc.close()
    return results

def blur_sensitive_pdf(pdf_path: str, output_path: str, language: str = 'en'):
    text_data = extract_text_from_pdf(pdf_path)
    doc = fitz.open(pdf_path)
    output_pdf = fitz.open()
    for page_number in range(len(doc)):
        page = doc[page_number]
        new_page = output_pdf.new_page(width=page.rect.width, height=page.rect.height)
        new_page.show_pdf_page(new_page.rect, doc, pno=page_number)
        page_text = text_data.get(page_number, "")
        if not page_text.strip():
            continue
        analysis = analyzer.analyze(text=page_text, language=language)
        for result in analysis:
            matched_text = page_text[result.start:result.end]
            bboxes = page.search_for(matched_text)
            for bbox in bboxes:
                new_page.draw_rect(
                    bbox,
                    color=(0.8, 0.8, 0.8),
                    fill=(0.8, 0.8, 0.8),
                    overlay=True
                )
                new_page.insert_text(
                    (bbox.x0 + 2, bbox.y0 + 10),
                    "[BLURRED]",
                    fontsize=6,
                    color=(0.5, 0.5, 0.5)
                )
    output_pdf.save(output_path)
    output_pdf.close()
    doc.close()

def process_pdf_file(pdf_path: str, output_dir: str = "output"):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    text_data = extract_text_from_pdf(pdf_path)
    doc = fitz.open(pdf_path)
    results = {}
    for page_num, text in text_data.items():
        if not text.strip():
            continue
        page = doc[int(page_num)]
        analysis = analyzer.analyze(text=text, language='en')
        sensitive_items = []
        for result in analysis:
            matched_text = text[result.start:result.end]
            bboxes = page.search_for(matched_text)
            for bbox in bboxes:
                sensitive_items.append({
                    'entity_type': result.entity_type,
                    'score': result.score,
                    'text': matched_text,
                    'bbox': [bbox.x0, bbox.y0, bbox.x1, bbox.y1]
                })
        results[str(page_num)] = sensitive_items
    doc.close()
    detected = results
    blurred_output = os.path.join(output_dir, f"blurred_{os.path.basename(pdf_path)}")
    blur_sensitive_pdf(pdf_path, blurred_output)
    encrypted_output = os.path.join(output_dir, f"encrypted_{os.path.basename(pdf_path)}")
    encrypt_sensitive_data(pdf_path, encrypted_output)
    return detected, blurred_output, encrypted_output

# --- DOCX ---
def extract_text_from_docx(docx_path: str) -> str:
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return ""

def detect_sensitive_docx(docx_path: str, language: str = 'en'):
    text_content = extract_text_from_docx(docx_path)
    if not text_content:
        return {'0': []}
    analysis = analyzer.analyze(text=text_content, language=language)
    sensitive_items = []
    for result in analysis:
        sensitive_items.append({
            'entity_type': result.entity_type,
            'score': result.score,
            'text': text_content[result.start:result.end],
            'start': result.start,
            'end': result.end
        })
    return {'0': sensitive_items}

def mask_docx_file(docx_path: str, output_path: str, language: str = 'en'):
    try:
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                analysis = analyzer.analyze(text=paragraph.text, language=language)
                if analysis:
                    masked_result = anonymizer.anonymize(text=paragraph.text, analyzer_results=analysis)
                    paragraph.text = masked_result.text
        doc.save(output_path)
    except Exception as e:
        print(f"Erreur lors du masquage DOCX: {e}")

def process_docx_file(docx_path: str, output_dir: str = "output"):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    text_content = extract_text_from_docx(docx_path)
    if not text_content:
        return {'0': []}, None, None
    analysis = analyzer.analyze(text=text_content, language='en')
    sensitive_items = []
    for result in analysis:
        sensitive_items.append({
            'entity_type': result.entity_type,
            'score': result.score,
            'text': text_content[result.start:result.end],
            'start': result.start,
            'end': result.end
        })
    detected = {'0': sensitive_items}
    masked_output = os.path.join(output_dir, f"masked_{os.path.basename(docx_path)}")
    mask_docx_file(docx_path, masked_output)
    encrypted_output = os.path.join(output_dir, f"encrypted_{os.path.basename(docx_path)}")
    encrypt_sensitive_data(docx_path, encrypted_output)
    return detected, masked_output, encrypted_output

# --- TXT ---
def extract_text_from_txt(txt_path: str) -> str:
    encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    for encoding in encodings_to_try:
        try:
            with open(txt_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            continue
    return ""

def detect_sensitive_txt(txt_path: str, language: str = 'en'):
    text_content = extract_text_from_txt(txt_path)
    if not text_content:
        return {'0': []}
    analysis = analyzer.analyze(text=text_content, language=language)
    sensitive_items = []
    for result in analysis:
        sensitive_items.append({
            'entity_type': result.entity_type,
            'score': result.score,
            'text': text_content[result.start:result.end],
            'start': result.start,
            'end': result.end
        })
    return {'0': sensitive_items}

def mask_txt_file(txt_path: str, output_path: str, language: str = 'en'):
    text_content = extract_text_from_txt(txt_path)
    if not text_content:
        return
    analysis = analyzer.analyze(text=text_content, language=language)
    masked_result = anonymizer.anonymize(text=text_content, analyzer_results=analysis)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(masked_result.text)

def process_txt_file(txt_path: str, output_dir: str = "output"):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    text_content = extract_text_from_txt(txt_path)
    if not text_content:
        return {'0': []}, None, None
    analysis = analyzer.analyze(text=text_content, language='en')
    sensitive_items = []
    for result in analysis:
        sensitive_items.append({
            'entity_type': result.entity_type,
            'score': result.score,
            'text': text_content[result.start:result.end],
            'start': result.start,
            'end': result.end
        })
    detected = {'0': sensitive_items}
    masked_output = os.path.join(output_dir, f"masked_{os.path.basename(txt_path)}")
    mask_txt_file(txt_path, masked_output)
    encrypted_output = os.path.join(output_dir, f"encrypted_{os.path.basename(txt_path)}")
    encrypt_sensitive_data(txt_path, encrypted_output)
    return detected, masked_output, encrypted_output

# --- XLSX ---
def extract_text_from_xlsx(xlsx_path: str) -> str:
    try:
        workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
        all_text = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            all_text.append(f"=== FEUILLE: {sheet_name} ===")
            for row in sheet.iter_rows():
                row_data = []
                for cell in row:
                    if cell.value is not None:
                        row_data.append(str(cell.value))
                    else:
                        row_data.append("")
                if any(cell_data.strip() for cell_data in row_data):
                    all_text.append("\t".join(row_data))
        workbook.close()
        return "\n".join(all_text)
    except Exception as e:
        return ""

def detect_sensitive_xlsx(xlsx_path: str, language: str = 'en'):
    text_content = extract_text_from_xlsx(xlsx_path)
    if not text_content:
        return {'0': []}
    analysis = analyzer.analyze(text=text_content, language=language)
    sensitive_items = []
    for result in analysis:
        sensitive_items.append({
            'entity_type': result.entity_type,
            'score': result.score,
            'text': text_content[result.start:result.end],
            'start': result.start,
            'end': result.end
        })
    return {'0': sensitive_items}

def mask_xlsx_file(xlsx_path: str, output_path: str, language: str = 'en'):
    try:
        workbook = openpyxl.load_workbook(xlsx_path)
        masked_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        cell_text = str(cell.value)
                        analysis = analyzer.analyze(text=cell_text, language=language)
                        if analysis:
                            masked_result = anonymizer.anonymize(text=cell_text, analyzer_results=analysis)
                            cell.value = masked_result.text
                            cell.fill = masked_fill
                            cell.comment = Comment("Données sensibles masquées", "System")
        workbook.save(output_path)
        workbook.close()
    except Exception as e:
        pass

def process_xlsx_file(xlsx_path: str, output_dir: str = "output"):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    text_content = extract_text_from_xlsx(xlsx_path)
    if not text_content:
        return {'0': []}, None, None
    analysis = analyzer.analyze(text=text_content, language='en')
    sensitive_items = []
    for result in analysis:
        sensitive_items.append({
            'entity_type': result.entity_type,
            'score': result.score,
            'text': text_content[result.start:result.end],
            'start': result.start,
            'end': result.end
        })
    detected = {'0': sensitive_items}
    masked_output = os.path.join(output_dir, f"masked_{os.path.basename(xlsx_path)}")
    mask_xlsx_file(xlsx_path, masked_output)
    encrypted_output = os.path.join(output_dir, f"encrypted_{os.path.basename(xlsx_path)}")
    encrypt_sensitive_data(xlsx_path, encrypted_output)
    return detected, masked_output, encrypted_output

# ----------- Chiffrement/déchiffrement TenSEAL -----------
def encrypt_sensitive_data(input_path: str, output_path: str, context_path: str = "tenseal_context.ctx"):
    ctx = get_ctx(context_path)
    file_type = detect_file_type(input_path)

    # Extraction selon type de fichier
    if file_type == 'csv':
        text_content = extract_text_from_csv(input_path)
        detection_results = detect_sensitive_csv(input_path)
    elif file_type == 'pdf':
        text_data = extract_text_from_pdf(input_path)
        detection_results = detect_sensitive_pdf(input_path)
        text_content = text_data  # dict
    elif file_type == 'docx':
        text_content = extract_text_from_docx(input_path)
        detection_results = detect_sensitive_docx(input_path)
    elif file_type == 'txt':
        text_content = extract_text_from_txt(input_path)
        detection_results = detect_sensitive_txt(input_path)
    elif file_type == 'xlsx':
        text_content = extract_text_from_xlsx(input_path)
        detection_results = detect_sensitive_xlsx(input_path)
    else:
        return

    ext = os.path.splitext(input_path)[1].lower()
    encrypted_data = []

    # ---- PDF : on stocke bbox, ciphertext ----
    if file_type == "pdf":
        doc = fitz.open(input_path)
        output_pdf = fitz.open()
        for page_number in range(len(doc)):
            page = doc[page_number]
            new_page = output_pdf.new_page(width=page.rect.width, height=page.rect.height)
            new_page.show_pdf_page(new_page.rect, doc, pno=page_number)
            page_str = str(page_number)
            if page_str not in detection_results:
                continue
            for item in detection_results[page_str]:
                pii_text = item['text']
                bbox_list = page.search_for(pii_text)
                for bbox in bbox_list:
                    plaintext_ints = [ord(c) for c in pii_text]
                    encrypted = ts.bfv_vector(ctx, plaintext_ints)
                    serialized = encrypted.serialize()
                    cipher_hex = serialized.hex()
                    encrypted_data.append({
                        "page": page_number,
                        "bbox": [float(bbox.x0), float(bbox.y0), float(bbox.x1), float(bbox.y1)],
                        "ciphertext": cipher_hex,
                        "length": len(plaintext_ints),
                        "text": pii_text
                    })
                    # Remplacement visible
                    new_page.draw_rect(bbox, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
                    new_page.insert_text((bbox.x0 + 2, bbox.y0 + 10), "<encrypted>", fontsize=6, color=(0, 0, 0))
        output_pdf.save(output_path)
        output_pdf.close()
        doc.close()
        with open(f"{output_path}_encrypted_data.json", "w") as f:
            json.dump(encrypted_data, f, indent=2)
    # ---- Autres fichiers texte ----
    else:
        text = text_content if isinstance(text_content, str) else "\n".join(text_content.values())
        replacements = []
        for item in detection_results['0']:
            pii_text = item['text']
            start, end = item['start'], item['end']
            plaintext_ints = [ord(c) for c in pii_text]
            encrypted = ts.bfv_vector(ctx, plaintext_ints)
            serialized = encrypted.serialize()
            cipher_hex = serialized.hex()
            encrypted_data.append({
                "start": start,
                "end": end,
                "ciphertext": cipher_hex,
                "length": len(plaintext_ints),
                "text": pii_text
            })
            replacements.append((start, end, "<encrypted>"))
        # Remplacement visible des PII
        replacements = sorted(replacements, key=lambda x: x[0], reverse=True)
        for start, end, repl in replacements:
            text = text[:start] + repl + text[end:]
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
        with open(f"{output_path}_encrypted_data.json", "w") as f:
            json.dump(encrypted_data, f, indent=2)

def decrypt_document(encrypted_path: str, output_path: str, context_path: str = "tenseal_context.ctx"):
    ctx = get_ctx(context_path)
    ext = os.path.splitext(encrypted_path)[1].lower()
    meta_path = f"{encrypted_path}_encrypted_data.json"

    # Check if metadata JSON exists
    if not os.path.exists(meta_path):
        return False

    # Load encrypted metadata
    with open(meta_path, "r") as f:
        encrypted_data = json.load(f)

    file_type = detect_file_type(encrypted_path)

    # --- PDF Decryption ---
    if file_type == "pdf":
        input_pdf = fitz.open(encrypted_path)
        output_pdf = fitz.open()

        for page_number in range(len(input_pdf)):
            page = input_pdf[page_number]
            new_page = output_pdf.new_page(width=page.rect.width, height=page.rect.height)
            new_page.show_pdf_page(new_page.rect, input_pdf, pno=page_number)

            for item in encrypted_data:
                if item["page"] != page_number:
                    continue

                bbox = item["bbox"]
                ciphertext = bytes.fromhex(item["ciphertext"])
                length = item["length"]

                encrypted_vec = ts.bfv_vector_from(ctx, ciphertext)
                decrypted = encrypted_vec.decrypt()

                # FIX: Map values back to 0–255
                decrypted_text = "".join(
                    [chr(int(x) % 256) for x in decrypted[:length] if 0 <= int(x) % 256 < 1114112]
                )

                bbox_rect = fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[3])
                new_page.draw_rect(bbox_rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
                new_page.insert_text((bbox[0] + 1, bbox[1] + 12), decrypted_text, fontsize=11, color=(0, 0, 0))

        output_pdf.save(output_path)
        output_pdf.close()
        input_pdf.close()
        return True

    # --- Text-based files (CSV, DOCX, TXT, XLSX) ---
    else:
        with open(encrypted_path, "r", encoding="utf-8") as f:
            text = f.read()

        replacements = []
        for item in encrypted_data:
            ciphertext = bytes.fromhex(item["ciphertext"])
            length = item["length"]

            encrypted_vec = ts.bfv_vector_from(ctx, ciphertext)
            decrypted = encrypted_vec.decrypt()

            # FIX: Map values back to 0–255
            decrypted_text = "".join(
                [chr(int(x) % 256) for x in decrypted[:length] if 0 <= int(x) % 256 < 1114112]
            )

            start, end = item["start"], item["end"]
            replacements.append((start, end, decrypted_text))

        # Apply replacements in reverse order
        replacements = sorted(replacements, key=lambda x: x[0], reverse=True)
        for start, end, repl in replacements:
            text = text[:start] + repl + text[start + len("<encrypted>"):]

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

        return True

# --- Dispatcher ---
def auto_detect_and_process(file_path: str, output_dir: str = "output"):
    file_type = detect_file_type(file_path)
    if file_type == 'csv':
        return process_csv_file(file_path, output_dir)
    elif file_type == 'pdf':
        return process_pdf_file(file_path, output_dir)
    elif file_type == 'docx':
        return process_docx_file(file_path, output_dir)
    elif file_type == 'txt':
        return process_txt_file(file_path, output_dir)
    elif file_type == 'xlsx':
        return process_xlsx_file(file_path, output_dir)
    else:
        return {}